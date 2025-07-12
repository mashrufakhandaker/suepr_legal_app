# app/services/doc_generate/doc_generate_service.py
from typing import Dict, Any
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches
import uuid

from app.services.shared.groq_client import GroqClient
from app.models.ai_models import LegalDocsInput, LegalDocsOutput

class DocGenerateService:
    def __init__(self):
        self.groq_client = GroqClient()
    
    async def generate_legal_document(self, input_data: LegalDocsInput) -> LegalDocsOutput:
        """Generate legal document based on input"""
        try:
            # Generate document content using AI
            system_prompt = f"""You are a legal document generator. Create a professional {input_data.document_type} 
            based on the provided information. The document should be:
            1. Professionally formatted
            2. Include all necessary legal language
            3. Be specific to the case details
            4. Include proper sender/recipient information
            5. Be actionable and clear
            
            Do not include placeholder text - use the actual information provided."""
            
            prompt = f"""
            Document Type: {input_data.document_type}
            Case Summary: {input_data.case_summary}
            
            Client Information:
            Name: {input_data.user_details.name}
            Address: {input_data.user_details.address}
            
            Opposing Party: {input_data.user_details.opposing_party}
            Case Facts: {input_data.user_details.facts}
            Additional Info: {input_data.user_details.additional_info or 'None'}
            
            Generate a complete, professional legal document.
            """
            
            doc_content = await self.groq_client.generate_response(
                prompt=prompt,
                system_prompt=system_prompt,
                max_tokens=2000
            )
            
            # Generate document title
            doc_title = f"{input_data.document_type} - {input_data.user_details.name}"
            
            # Create Word document
            doc_path = await self._create_word_document(doc_content, doc_title, input_data.user_details)
            
            return LegalDocsOutput(
                doc_title=doc_title,
                doc_content=doc_content,
                format="docx",
                download_url=f"/uploads/{os.path.basename(doc_path)}"
            )
            
        except Exception as e:
            raise Exception(f"Error generating legal document: {str(e)}")
    
    async def _create_word_document(self, content: str, title: str, user_details) -> str:
        """Create a Word document from the generated content"""
        try:
            # Create new document
            doc = Document()
            
            # Add title
            title_para = doc.add_heading(title, 0)
            title_para.alignment = 1  # Center alignment
            
            # Add date
            date_para = doc.add_paragraph(f"Date: {datetime.now().strftime('%B %d, %Y')}")
            date_para.alignment = 2  # Right alignment
            
            # Add sender information
            doc.add_paragraph()
            sender_para = doc.add_paragraph("From:")
            sender_para.add_run(f"\n{user_details.name}")
            sender_para.add_run(f"\n{user_details.address}")
            
            # Add recipient information
            doc.add_paragraph()
            recipient_para = doc.add_paragraph("To:")
            recipient_para.add_run(f"\n{user_details.opposing_party}")
            
            # Add main content
            doc.add_paragraph()
            doc.add_paragraph(content)
            
            # Add signature block
            doc.add_paragraph()
            doc.add_paragraph("Sincerely,")
            doc.add_paragraph()
            doc.add_paragraph("_" * 30)
            doc.add_paragraph(user_details.name)
            
            # Save document
            filename = f"{title.replace(' ', '_')}_{uuid.uuid4().hex[:8]}.docx"
            filepath = os.path.join("uploads", filename)
            doc.save(filepath)
            
            return filepath
            
        except Exception as e:
            raise Exception(f"Error creating Word document: {str(e)}")
    
    async def get_document_templates(self) -> Dict[str, Any]:
        """Get available document templates"""
        templates = {
            "Demand Letter": {
                "description": "Formal demand for payment or action",
                "required_fields": ["opposing_party", "amount_owed", "deadline"],
                "typical_use": "Debt collection, contract disputes"
            },
            "Cease and Desist": {
                "description": "Request to stop specific behavior",
                "required_fields": ["opposing_party", "behavior_to_stop", "legal_basis"],
                "typical_use": "Harassment, copyright infringement"
            },
            "Notice to Quit": {
                "description": "Eviction notice for tenants",
                "required_fields": ["tenant_name", "property_address", "violation_reason"],
                "typical_use": "Landlord-tenant disputes"
            },
            "Small Claims Petition": {
                "description": "Filing for small claims court",
                "required_fields": ["defendant_name", "claim_amount", "claim_basis"],
                "typical_use": "Small monetary disputes"
            },
            "Contract": {
                "description": "Basic service or agreement contract",
                "required_fields": ["parties", "terms", "payment_details"],
                "typical_use": "Service agreements, sales contracts"
            }
        }
        
        return templates
